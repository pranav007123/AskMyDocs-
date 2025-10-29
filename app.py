"""
Flask RAG Application with User Authentication
"""
import os
import uuid
import shutil
from datetime import datetime
from typing import List, Dict, Any

import openai
try:
    import google.generativeai as genai
except Exception:
    genai = None
import fitz  # PyMuPDF
import docx
import faiss
import numpy as np
from sentence_transformers import SentenceTransformer
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, abort
from werkzeug.utils import secure_filename
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from functools import wraps
from flask_migrate import Migrate
from dotenv import load_dotenv

from models import db, User, Document

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'dev-secret-key-change-in-production')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///db.sqlite3'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
app.config['UPLOAD_FOLDER'] = os.path.join(BASE_DIR, 'uploads')
app.config['FAISS_FOLDER'] = os.path.join(BASE_DIR, 'faiss_indexes')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Initialize extensions
db.init_app(app)
migrate = Migrate(app, db)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Please log in to access this page.'

# LLM provider configuration
LLM_PROVIDER = os.getenv('LLM_PROVIDER', 'openai').lower().strip()  # 'openai' (default) or 'gemini'
print(f"[STARTUP] LLM_PROVIDER: {LLM_PROVIDER}")

# Initialize OpenAI or Gemini
openai.api_key = os.getenv('OPENAI_API_KEY')
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
print(f"[STARTUP] GEMINI_API_KEY configured: {bool(GEMINI_API_KEY)}")
print(f"[STARTUP] OPENAI_API_KEY configured: {bool(openai.api_key)}")

if LLM_PROVIDER == 'gemini' and genai is not None and GEMINI_API_KEY:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        print(f"[STARTUP] Gemini client configured successfully")
    except Exception as e:
        print(f"Warning: Failed to configure Gemini client: {e}")
elif LLM_PROVIDER == 'gemini':
    print(f"[STARTUP] Warning: Gemini selected but not properly configured")

# Initialize sentence transformer model
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')

# Allowed file extensions
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}

@login_manager.user_loader
def load_user(user_id):
    """Load user for Flask-Login."""
    return User.query.get(int(user_id))

def admin_required(f):
    """Decorator to require admin privileges."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated:
            flash('Please log in to access this page.', 'error')
            return redirect(url_for('login'))
        if not current_user.is_admin:
            flash('You do not have permission to access this page.', 'error')
            return redirect(url_for('dashboard'))
        return f(*args, **kwargs)
    return decorated_function

def validate_password(password):
    """Validate password strength.
    Returns (is_valid, error_message)
    """
    if len(password) < 8:
        return False, "Password must be at least 8 characters long."
    
    if not any(char.isupper() for char in password):
        return False, "Password must contain at least one uppercase letter."
    
    if not any(char.islower() for char in password):
        return False, "Password must contain at least one lowercase letter."
    
    if not any(char.isdigit() for char in password):
        return False, "Password must contain at least one digit."
    
    # Optional: Check for special characters
    special_chars = "!@#$%^&*()_+-=[]{}|;:,.<>?"
    if not any(char in special_chars for char in password):
        return False, "Password must contain at least one special character (!@#$%^&* etc.)."
    
    return True, "Password is strong."

def allowed_file(filename):
    """Check if file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def create_directories():
    """Create necessary directories if they don't exist."""
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(app.config['FAISS_FOLDER'], exist_ok=True)

def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract text from uploaded files."""
    try:
        if file_type == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif file_type == 'pdf':
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        
        elif file_type == 'docx':
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        return ""
    except Exception as e:
        print(f"Error extracting text from {file_path}: {str(e)}")
        return ""

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 100) -> List[str]:
    """Split text into overlapping chunks."""
    words = text.split()
    chunks = []
    
    for i in range(0, len(words), chunk_size - overlap):
        chunk = ' '.join(words[i:i + chunk_size])
        if chunk.strip():
            chunks.append(chunk.strip())
        
        if i + chunk_size >= len(words):
            break
    
    return chunks

def get_user_faiss_path(user_id: int) -> str:
    """Get the FAISS index path for a user."""
    user_dir = os.path.join(app.config['FAISS_FOLDER'], str(user_id))
    os.makedirs(user_dir, exist_ok=True)
    return os.path.join(user_dir, 'index.faiss')

def get_user_metadata_path(user_id: int) -> str:
    """Get the metadata path for a user."""
    user_dir = os.path.join(app.config['FAISS_FOLDER'], str(user_id))
    os.makedirs(user_dir, exist_ok=True)
    return os.path.join(user_dir, 'metadata.npy')

def load_or_create_faiss_index(user_id: int):
    """Load existing FAISS index or create a new one."""
    faiss_path = get_user_faiss_path(user_id)
    metadata_path = get_user_metadata_path(user_id)
    
    if os.path.exists(faiss_path) and os.path.exists(metadata_path):
        index = faiss.read_index(faiss_path)
        metadata = np.load(metadata_path, allow_pickle=True).tolist()
        return index, metadata
    else:
        # Create new index (384 dimensions for all-MiniLM-L6-v2)
        index = faiss.IndexFlatIP(384)
        metadata = []
        return index, metadata

def save_faiss_index(user_id: int, index, metadata):
    """Save FAISS index and metadata."""
    faiss_path = get_user_faiss_path(user_id)
    metadata_path = get_user_metadata_path(user_id)
    
    faiss.write_index(index, faiss_path)
    np.save(metadata_path, metadata)

def add_document_to_faiss(user_id: int, document_id: int, chunks: List[str], filename: str):
    """Add document chunks to user's FAISS index."""
    index, metadata = load_or_create_faiss_index(user_id)
    
    # Generate embeddings for chunks
    embeddings = embedding_model.encode(chunks)
    
    # Normalize embeddings for cosine similarity
    faiss.normalize_L2(embeddings)
    
    # Add to index
    index.add(embeddings)
    
    # Add metadata
    for i, chunk in enumerate(chunks):
        metadata.append({
            'document_id': document_id,
            'chunk_index': i,
            'text': chunk,
            'filename': filename
        })
    
    # Save updated index
    save_faiss_index(user_id, index, metadata)
    
    return len(chunks)

def highlight_relevant_content(text: str, query: str) -> str:
    """Highlight content most relevant to the query using semantic similarity."""
    import re
    from sentence_transformers import util
    
    # Split text into sentences
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if s.strip()]
    
    if not sentences or not query.strip():
        return text
    
    try:
        # Generate embeddings for query and sentences
        query_embedding = embedding_model.encode([query])
        sentence_embeddings = embedding_model.encode(sentences)
        
        # Calculate similarity scores
        similarities = util.cos_sim(query_embedding, sentence_embeddings)[0]
        
        # Find sentences with high similarity (top 30% or score > 0.3)
        threshold = max(0.3, similarities.mean() + similarities.std() * 0.5)
        relevant_sentences = []
        
        for i, (sentence, score) in enumerate(zip(sentences, similarities)):
            if score > threshold:
                relevant_sentences.append(sentence)
        
        # If no sentences meet threshold, take top 2 most similar
        if not relevant_sentences:
            top_indices = similarities.argsort(descending=True)[:2]
            relevant_sentences = [sentences[i] for i in top_indices if i < len(sentences)]
        
        # Highlight relevant sentences in the original text
        highlighted_text = text
        for sentence in relevant_sentences:
            if sentence in highlighted_text:
                highlighted_sentence = f'<mark class="search-highlight">{sentence}</mark>'
                highlighted_text = highlighted_text.replace(sentence, highlighted_sentence)
        
        return highlighted_text
        
    except Exception as e:
        print(f"Error in semantic highlighting: {str(e)}")
        # Fallback to keyword highlighting
        return highlight_keywords(text, query)

def highlight_keywords(text: str, query: str) -> str:
    """Fallback keyword highlighting."""
    import re
    
    # Extract key terms from query (nouns, verbs, important words)
    query_words = [word.strip().lower() for word in re.split(r'[^\w]+', query) 
                   if word.strip() and len(word) > 2]
    
    # Filter out common stop words
    stop_words = {'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'what', 'how', 'why', 'when', 'where'}
    query_words = [word for word in query_words if word not in stop_words]
    
    if not query_words:
        return text
    
    # Create pattern for important words
    pattern = r'\b(' + '|'.join(re.escape(word) for word in query_words) + r')\b'
    
    def replace_match(match):
        return f'<mark class="search-highlight">{match.group(0)}</mark>'
    
    return re.sub(pattern, replace_match, text, flags=re.IGNORECASE)

def search_faiss_index(user_id: int, query: str, k: int = 5, document_id: int = None):
    """Search FAISS index for relevant chunks."""
    try:
        index, metadata = load_or_create_faiss_index(user_id)
        
        if index.ntotal == 0:
            return []
        
        # Generate query embedding
        query_embedding = embedding_model.encode([query])
        faiss.normalize_L2(query_embedding)
        
        # Search
        scores, indices = index.search(query_embedding, min(k, index.ntotal))
        
        results = []
        for score, idx in zip(scores[0], indices[0]):
            if idx < len(metadata):
                chunk_metadata = metadata[idx]
                
                # Filter by document if specified
                if document_id and chunk_metadata['document_id'] != document_id:
                    continue
                
                # Highlight semantically relevant content
                highlighted_text = highlight_relevant_content(chunk_metadata['text'], query)
                
                results.append({
                    'text': chunk_metadata['text'],  # Original text for LLM
                    'highlighted_text': highlighted_text,  # Highlighted text for display
                    'filename': chunk_metadata['filename'],
                    'document_id': chunk_metadata['document_id'],
                    'score': float(score)
                })
        
        return results
    except Exception as e:
        print(f"Error searching FAISS index: {str(e)}")
        return []

def _build_context_and_prompt(query: str, context_chunks: List[Dict]) -> str:
    context = "\n\n".join([
        f"From {chunk['filename']}:\n{chunk['text']}"
        for chunk in context_chunks
    ])
    prompt = (
        "Based on the following context from the user's documents, please answer the question. "
        "If the context doesn't contain enough information to answer the question, please say so.\n\n"
        f"Context:\n{context}\n\nQuestion: {query}\n\nAnswer:"
    )
    return prompt

def generate_rag_response_openai(query: str, context_chunks: List[Dict]) -> str:
    if not context_chunks:
        return "I couldn't find any relevant information in your documents to answer this question."
    prompt = _build_context_and_prompt(query, context_chunks)
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that answers questions based on provided context from documents. Be accurate and cite the source documents when possible."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=500,
            temperature=0.7
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"Error generating response (OpenAI): {str(e)}")
        return f"Sorry, I encountered an error while generating the response: {str(e)}"

def generate_rag_response_gemini(query: str, context_chunks: List[Dict]) -> str:
    if not context_chunks:
        return "I couldn't find any relevant information in your documents to answer this question."
    if genai is None:
        return "Gemini backend not available on this server."
    if not GEMINI_API_KEY:
        return "Gemini API key not configured. Please set GEMINI_API_KEY in .env."
    prompt = _build_context_and_prompt(query, context_chunks)
    # Try multiple candidate models for better compatibility
    candidates = choose_gemini_model(candidates_only=True)
    last_err = None
    for model_name in candidates:
        try:
            print(f"[RAG] Trying Gemini model: {model_name}")
            model = genai.GenerativeModel(model_name)
            resp = model.generate_content(prompt)
            return (getattr(resp, 'text', '') or '').strip() or str(resp)
        except Exception as e:
            last_err = e
            print(f"[Gemini] Model '{model_name}' failed: {e}")
            # If rate-limited by Gemini and OpenAI is available, fall back to OpenAI automatically
            err_text = str(e)
            if ('429' in err_text or 'rate' in err_text.lower()) and (openai.api_key):
                print("[RAG] Gemini rate-limited. Falling back to OpenAI...")
                try:
                    return generate_rag_response_openai(query, context_chunks)
                except Exception as openai_err:
                    print(f"[OpenAI Fallback] Error: {openai_err}")
                    # If fallback also fails, continue trying other Gemini candidates
            continue
    # If all candidates failed, surface the last error
    err_text = str(last_err)
    if ('429' in err_text or 'rate' in err_text.lower()) and not openai.api_key:
        return (
            "Gemini is currently rate-limited for this project. "
            "Add OPENAI_API_KEY in .env and set LLM_PROVIDER=openai to continue immediately, "
            "or retry after a short wait."
        )
    return f"Sorry, I encountered an error while generating the response: {err_text}"

def choose_gemini_model(candidates_only: bool = False):
    """Pick an available Gemini model that supports text generation.
    Strategy:
    1) Prefer a supported model from the API that includes 'generateContent'.
    2) Fallback through a list of known good model IDs.
    """
    try:
        # Prefer models that support text generation
        available = []
        try:
            models = list(genai.list_models())
            for m in models:
                methods = set(getattr(m, 'supported_generation_methods', []) or [])
                if 'generateContent' in methods or 'generate_content' in methods:
                    available.append(getattr(m, 'name', ''))
        except Exception as e:
            print(f"[Gemini] Could not list models: {e}")

        # Preferred order: latest first
        preferred = [
            'gemini-1.5-flash-latest',
            'gemini-1.5-pro-latest',
            'gemini-1.5-flash-8b',
            'gemini-1.5-flash',
            'gemini-1.5-pro',
            'gemini-1.0-pro',
        ]

        for name in preferred:
            if name in available:
                return preferred if candidates_only else name

        # If listing failed or names differ (some APIs return fully qualified names)
        # try constructing from fully qualified names too
        fq_preferred = [
            'models/gemini-1.5-flash-latest',
            'models/gemini-1.5-pro-latest',
            'models/gemini-1.5-flash',
            'models/gemini-1.5-flash-8b',
            'models/gemini-1.5-pro',
            'models/gemini-1.0-pro',
        ]
        for name in fq_preferred:
            if name in available:
                return (preferred + fq_preferred) if candidates_only else name

        # Final fallback: return all candidates: preferred + fq_preferred + first available
        if available:
            return (preferred + fq_preferred + available) if candidates_only else available[0]
    except Exception as e:
        print(f"[Gemini] Model selection error: {e}")

    # If everything fails, return sensible candidates
    fallback_list = [
        'gemini-1.5-flash-latest',
        'gemini-1.5-pro-latest',
        'models/gemini-1.5-pro',
        'models/gemini-1.5-flash',
        'gemini-1.0-pro',
    ]
    return fallback_list if candidates_only else fallback_list[0]

def generate_rag_response(query: str, context_chunks: List[Dict]) -> str:
    """Generate response using the selected LLM provider."""
    print(f"[RAG] Using LLM provider: {LLM_PROVIDER}")
    if LLM_PROVIDER == 'gemini':
        print(f"[RAG] Calling Gemini API")
        return generate_rag_response_gemini(query, context_chunks)
    print(f"[RAG] Calling OpenAI API")
    return generate_rag_response_openai(query, context_chunks)

def remove_document_from_faiss(user_id: int, document_id: int):
    """Remove document chunks from FAISS index."""
    try:
        index, metadata = load_or_create_faiss_index(user_id)
        
        # Find indices to remove
        indices_to_remove = []
        new_metadata = []
        
        for i, chunk_meta in enumerate(metadata):
            if chunk_meta['document_id'] == document_id:
                indices_to_remove.append(i)
            else:
                new_metadata.append(chunk_meta)
        
        if indices_to_remove:
            # Create new index without the removed chunks
            if new_metadata:
                # Get embeddings for remaining chunks
                remaining_texts = [meta['text'] for meta in new_metadata]
                embeddings = embedding_model.encode(remaining_texts)
                faiss.normalize_L2(embeddings)
                
                # Create new index
                new_index = faiss.IndexFlatIP(384)
                new_index.add(embeddings)
                
                save_faiss_index(user_id, new_index, new_metadata)
            else:
                # No chunks left, create empty index
                empty_index = faiss.IndexFlatIP(384)
                save_faiss_index(user_id, empty_index, [])
    
    except Exception as e:
        print(f"Error removing document from FAISS: {str(e)}")

# Routes
@app.route('/')
def index():
    """Home page."""
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return render_template('index.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    """User registration."""
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '')
        confirm_password = request.form.get('confirm_password', '')
        
        # Validate inputs
        if not username or not email or not password:
            flash('All fields are required.', 'error')
            return render_template('register.html')
        
        # Validate username length
        if len(username) < 3:
            flash('Username must be at least 3 characters long.', 'error')
            return render_template('register.html')
        
        # Validate email format
        import re
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(email_pattern, email):
            flash('Please enter a valid email address.', 'error')
            return render_template('register.html')
        
        # Validate password confirmation
        if confirm_password and password != confirm_password:
            flash('Passwords do not match.', 'error')
            return render_template('register.html')
        
        # Validate password strength
        is_valid, message = validate_password(password)
        if not is_valid:
            flash(message, 'error')
            return render_template('register.html')
        
        # Check if username exists
        if User.query.filter_by(username=username).first():
            flash('Username already exists.', 'error')
            return render_template('register.html')
        
        # Check if email exists
        if User.query.filter_by(email=email).first():
            flash('Email already registered.', 'error')
            return render_template('register.html')
        
        # Create new user
        user = User(username=username, email=email)
        user.set_password(password)
        
        db.session.add(user)
        db.session.commit()
        
        flash('Registration successful! Please log in.', 'success')
        return redirect(url_for('login'))
    
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    """User login."""
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        
        user = User.query.filter_by(email=email).first()
        
        if user and user.check_password(password):
            login_user(user)
            flash('Logged in successfully!', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid email or password.', 'error')
    
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    """User logout."""
    logout_user()
    flash('Logged out successfully!', 'success')
    return redirect(url_for('index'))

@app.route('/dashboard')
@login_required
def dashboard():
    """User dashboard with document management."""
    documents = Document.query.filter_by(user_id=current_user.id).order_by(Document.uploaded_at.desc()).all()
    return render_template('dashboard.html', documents=documents)

@app.route('/upload', methods=['POST'])
@login_required
def upload_file():
    """Handle file upload."""
    try:
        app.logger.info('Upload attempt by user_id=%s', getattr(current_user, 'id', None))
        print(f"[UPLOAD] Request files: {list(request.files.keys())}")
        print(f"[UPLOAD] Request form: {dict(request.form)}")
        
        if 'file' not in request.files:
            flash('No file selected.', 'error')
            app.logger.warning('Upload failed: no file part in request')
            return redirect(url_for('dashboard'))
        
        file = request.files['file']
        print(f"[UPLOAD] File object: {file}")
        print(f"[UPLOAD] Filename: {file.filename}")
        
        if file.filename == '':
            flash('No file selected.', 'error')
            app.logger.warning('Upload failed: empty filename')
            return redirect(url_for('dashboard'))
    except Exception as e:
        print(f"[UPLOAD] Error in initial checks: {str(e)}")
        flash(f'Upload error: {str(e)}', 'error')
        return redirect(url_for('dashboard'))
    
    if file and allowed_file(file.filename):
        try:
            print(f"[UPLOAD] Processing file: {file.filename}")
            # Generate unique filename
            original_filename = secure_filename(file.filename)
            file_extension = original_filename.rsplit('.', 1)[1].lower()
            unique_filename = f"{uuid.uuid4()}.{file_extension}"
            
            print(f"[UPLOAD] Original: {original_filename}, Unique: {unique_filename}")
            
            # Save file
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
            print(f"[UPLOAD] Saving to: {file_path}")
        except Exception as e:
            print(f"[UPLOAD] Error in file processing setup: {str(e)}")
            flash(f'Error processing file: {str(e)}', 'error')
            return redirect(url_for('dashboard'))
        app.logger.info('Saving upload to %s', file_path)
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        file.save(file_path)
        app.logger.info('Saved file: %s (size=%s bytes)', file_path, os.path.getsize(file_path) if os.path.exists(file_path) else 'N/A')
        
        try:
            print(f"[UPLOAD] Extracting text from {file_extension} file...")
            # Extract text
            text = extract_text_from_file(file_path, file_extension)
            print(f"[UPLOAD] Extracted {len(text)} characters")
            
            if not text.strip():
                flash('Could not extract text from the file. Please ensure it contains readable text.', 'error')
                app.logger.warning('Upload failed: no text extracted from %s', original_filename)
                if os.path.exists(file_path):
                    os.remove(file_path)
                return redirect(url_for('dashboard'))
            
            # Create document record
            document = Document(
                user_id=current_user.id,
                filename=unique_filename,
                original_name=original_filename,
                file_type=file_extension
            )
            db.session.add(document)
            db.session.commit()
            app.logger.info('Created Document id=%s for user_id=%s', document.id, current_user.id)
            
            # Process text into chunks
            chunks = chunk_text(text)
            app.logger.info('Chunked text into %s chunks', len(chunks))
            
            # Add to FAISS index
            chunk_count = add_document_to_faiss(
                current_user.id, 
                document.id, 
                chunks, 
                original_filename
            )
            
            # Update chunk count
            document.chunk_count = chunk_count
            db.session.commit()
            app.logger.info('Indexed document id=%s with %s chunks', document.id, chunk_count)
            
            flash(f'File "{original_filename}" uploaded successfully! Created {chunk_count} chunks.', 'success')
        
        except Exception as e:
            # Clean up on error
            if os.path.exists(file_path):
                os.remove(file_path)
            if 'document' in locals():
                db.session.delete(document)
                db.session.commit()
            
            app.logger.exception('Error processing file %s: %s', original_filename, str(e))
            flash(f'Error processing file: {str(e)}', 'error')
    
    else:
        flash('Invalid file type. Please upload PDF, DOCX, or TXT files.', 'error')
        app.logger.warning('Upload failed: invalid file type: %s', file.filename)
    
    return redirect(url_for('dashboard'))

@app.errorhandler(413)
def request_entity_too_large(error):
    """Handle files exceeding MAX_CONTENT_LENGTH."""
    flash('File too large. Maximum allowed size is 16MB.', 'error')
    app.logger.warning('Upload rejected: file too large')
    return redirect(url_for('dashboard'))

@app.route('/delete_document/<int:document_id>')
@login_required
def delete_document(document_id):
    """Delete a document."""
    document = Document.query.filter_by(id=document_id, user_id=current_user.id).first()
    
    if not document:
        flash('Document not found.', 'error')
        return redirect(url_for('dashboard'))
    
    try:
        # Remove from FAISS index
        remove_document_from_faiss(current_user.id, document_id)
        
        # Remove file
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], document.filename)
        if os.path.exists(file_path):
            os.remove(file_path)
        
        # Remove from database
        db.session.delete(document)
        db.session.commit()
        
        flash(f'Document "{document.original_name}" deleted successfully.', 'success')
    
    except Exception as e:
        flash(f'Error deleting document: {str(e)}', 'error')
    
    return redirect(url_for('dashboard'))

@app.route('/search')
@login_required
def search_page():
    """Search page."""
    documents = Document.query.filter_by(user_id=current_user.id).all()
    return render_template('search.html', documents=documents)

@app.route('/search', methods=['POST'])
@login_required
def search():
    """Handle search queries."""
    query = request.form.get('query', '').strip()
    document_id = request.form.get('document_id')
    
    print(f"[SEARCH] User {current_user.id} searching for: '{query}'")
    
    if not query:
        flash('Please enter a search query.', 'error')
        return redirect(url_for('search_page'))
    
    try:
        # Check if user has any documents
        user_documents = Document.query.filter_by(user_id=current_user.id).all()
        print(f"[SEARCH] User has {len(user_documents)} documents")
        
        if not user_documents:
            flash('You need to upload documents before searching. Please upload some documents first.', 'warning')
            return redirect(url_for('dashboard'))
        
        # Convert document_id to int if provided
        doc_id = int(document_id) if document_id and document_id != 'all' else None
        print(f"[SEARCH] Searching in document_id: {doc_id}")
        
        # Search FAISS index
        results = search_faiss_index(current_user.id, query, k=5, document_id=doc_id)
        print(f"[SEARCH] Found {len(results)} results")
        
        if not results:
            # Check if FAISS index exists and has content
            index, metadata = load_or_create_faiss_index(current_user.id)
            print(f"[SEARCH] FAISS index has {index.ntotal} vectors, {len(metadata)} metadata entries")
            
            if index.ntotal == 0:
                flash('Your documents are still being processed. Please try again in a moment, or re-upload your documents.', 'warning')
            else:
                flash('No relevant information found for your query. Try rephrasing your question or using different keywords.', 'info')
            
            return render_template('search.html',
                                 documents=user_documents,
                                 query=query)
        
        # Generate RAG response
        print(f"[SEARCH] Generating RAG response...")
        answer = generate_rag_response(query, results)
        print(f"[SEARCH] Generated answer: {len(answer)} characters")
        
        return render_template('search.html',
                             documents=user_documents,
                             query=query,
                             answer=answer,
                             sources=results)
    
    except Exception as e:
        print(f"[SEARCH] Error: {str(e)}")
        import traceback
        traceback.print_exc()
        flash(f'Error processing search: {str(e)}', 'error')
        return redirect(url_for('search_page'))

@app.route('/download/<int:document_id>')
@login_required
def download_document(document_id):
    """Download a document."""
    document = Document.query.filter_by(id=document_id, user_id=current_user.id).first()
    
    if not document:
        flash('Document not found.', 'error')
        return redirect(url_for('dashboard'))
    
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], document.filename)
    
    if not os.path.exists(file_path):
        flash('File not found on server.', 'error')
        return redirect(url_for('dashboard'))
    
    return send_file(file_path, as_attachment=True, download_name=document.original_name)

# Admin Routes
@app.route('/admin')
@login_required
@admin_required
def admin_dashboard():
    """Admin dashboard with analytics."""
    # Get statistics
    total_users = User.query.count()
    total_documents = Document.query.count()
    total_chunks = db.session.query(db.func.sum(Document.chunk_count)).scalar() or 0
    
    # Recent users (last 10)
    recent_users = User.query.order_by(User.created_at.desc()).limit(10).all()
    
    # Recent documents (last 10)
    recent_documents = Document.query.order_by(Document.uploaded_at.desc()).limit(10).all()
    
    # User statistics
    users_with_docs = db.session.query(User).join(Document).distinct().count()
    users_without_docs = total_users - users_with_docs
    
    # Document statistics by type
    doc_stats = db.session.query(
        Document.file_type,
        db.func.count(Document.id).label('count')
    ).group_by(Document.file_type).all()
    
    # Top users by document count
    top_users = db.session.query(
        User,
        db.func.count(Document.id).label('doc_count')
    ).join(Document).group_by(User.id).order_by(db.desc('doc_count')).limit(5).all()
    
    return render_template('admin_dashboard.html',
                         total_users=total_users,
                         total_documents=total_documents,
                         total_chunks=total_chunks,
                         recent_users=recent_users,
                         recent_documents=recent_documents,
                         users_with_docs=users_with_docs,
                         users_without_docs=users_without_docs,
                         doc_stats=doc_stats,
                         top_users=top_users)

@app.route('/admin/users')
@login_required
@admin_required
def admin_users():
    """Admin page to view all users."""
    users = User.query.order_by(User.created_at.desc()).all()
    
    # Get document count for each user
    user_stats = []
    for user in users:
        doc_count = Document.query.filter_by(user_id=user.id).count()
        chunk_count = db.session.query(db.func.sum(Document.chunk_count)).filter(
            Document.user_id == user.id
        ).scalar() or 0
        user_stats.append({
            'user': user,
            'doc_count': doc_count,
            'chunk_count': chunk_count
        })
    
    return render_template('admin_users.html', user_stats=user_stats)

@app.route('/admin/user/<int:user_id>/toggle_admin', methods=['POST'])
@login_required
@admin_required
def toggle_admin(user_id):
    """Toggle admin status for a user."""
    if user_id == current_user.id:
        flash('You cannot change your own admin status.', 'error')
        return redirect(url_for('admin_users'))
    
    user = User.query.get_or_404(user_id)
    user.is_admin = not user.is_admin
    db.session.commit()
    
    status = 'admin' if user.is_admin else 'regular user'
    flash(f'User {user.username} is now a {status}.', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/user/<int:user_id>/delete', methods=['POST'])
@login_required
@admin_required
def admin_delete_user(user_id):
    """Delete a user and all their documents."""
    if user_id == current_user.id:
        flash('You cannot delete your own account from admin panel.', 'error')
        return redirect(url_for('admin_users'))
    
    user = User.query.get_or_404(user_id)
    username = user.username
    
    try:
        # Delete all user's documents and FAISS indexes
        documents = Document.query.filter_by(user_id=user_id).all()
        for doc in documents:
            # Remove file
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], doc.filename)
            if os.path.exists(file_path):
                os.remove(file_path)
        
        # Remove FAISS index directory
        user_faiss_dir = os.path.join(app.config['FAISS_FOLDER'], str(user_id))
        if os.path.exists(user_faiss_dir):
            shutil.rmtree(user_faiss_dir)
        
        # Delete user (cascade will delete documents)
        db.session.delete(user)
        db.session.commit()
        
        flash(f'User {username} and all associated data deleted successfully.', 'success')
    except Exception as e:
        flash(f'Error deleting user: {str(e)}', 'error')
    
    return redirect(url_for('admin_users'))

if __name__ == '__main__':
    create_directories()
    
    with app.app_context():
        db.create_all()
    
    app.run(debug=True, host='0.0.0.0', port=5000)
